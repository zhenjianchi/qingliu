#!/usr/bin/env python3
"""生成戒除成瘾工具 PRD 文档"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ===== 页面设置 =====
section = doc.sections[0]
section.page_width = Inches(11.69)   # A4
section.page_height = Inches(16.54)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.18)
section.right_margin = Inches(1.18)

# ===== 字体 =====
def set_run_font(run, name='Calibri', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 中文字体
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1, color=(30, 58, 95)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        set_run_font(run, 'Calibri', 16, bold=True, color=color)
    elif level == 2:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_run_font(run, 'Calibri', 13, bold=True, color=color)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_run_font(run, 'Calibri', 11, bold=True, color=color)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_run_font(run, 'Calibri', 11)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    set_run_font(run, 'Calibri', 11)
    return p

def add_numbered(doc, text, num, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3 + indent * 0.25)
    run_num = p.add_run(f"{num}. ")
    set_run_font(run_num, 'Calibri', 11, bold=True, color=(30, 58, 95))
    run_text = p.add_run(text)
    set_run_font(run_text, 'Calibri', 11)
    return p

# ===== 封面 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
run = p.add_run("清 流")
set_run_font(run, 'Calibri', 36, bold=True, color=(30, 58, 95))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("健康习惯管理工具")
set_run_font(run, 'Calibri', 18, color=(68, 114, 196))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("产品需求文档（PRD）")
set_run_font(run, 'Calibri', 14, bold=True)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("v1.0 · MVP 版本")
set_run_font(run, 'Calibri', 12, color=(89, 89, 89))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"编制日期：{datetime.date.today().strftime('%Y年%m月%d日')}")
set_run_font(run, 'Calibri', 11, color=(89, 89, 89))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("密级：内部 · 仅供团队使用")
set_run_font(run, 'Calibri', 10, color=(150, 150, 150))

doc.add_page_break()

# ===== 正文 =====

# 1. 文档概述
add_heading(doc, "一、文档概述", 1)
add_body(doc, "本文档为「清流」健康习惯管理工具的 MVP（最小可行产品）版本产品需求文档，定义了离线工具形态下需要实现的全部核心功能、交互逻辑和技术规格。")

add_heading(doc, "1.1 产品定位", 2)
add_body(doc, "「清流」是一款帮助用户戒除色情内容沉迷与手淫等强迫性行为的健康习惯管理工具，采用离线优先设计，隐私高度保护，界面完全中性化（可伪装为普通健康类应用）。")

add_body(doc, "核心价值主张：在用户最脆弱的渴望高峰时刻，提供 3 秒内可执行的神经打断技术，将本能冲动转化为可管理的短暂不适，从而打破行为强化回路。")

add_heading(doc, "1.2 目标用户", 2)
bullets = [
    "15-35 岁、有戒除色情/手淫成瘾意愿但缺乏有效工具的青少年和青年",
    "追求更健康生活方式、关注身心平衡的用户",
    "（非当前 MVP 范围）希望帮助伴侣/子女建立健康习惯的家庭成员",
]
for b in bullets:
    add_bullet(doc, b)

add_heading(doc, "1.3 设计理念", 2)
principles = [
    ("隐私优先", "数据完全本地存储，不强制云端；应用名称和界面中性化"),
    ("神经科学驱动", "所有渴望应对技术均来自经科学验证的神经打断机制"),
    ("正向激励", "强调进步而非惩罚，relapse 后不批判，引导继续前行"),
    ("极简上手", "3 分钟内可开始使用，核心流程无学习成本"),
]
for name, desc in principles:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"● {name}：")
    set_run_font(run, 'Calibri', 11, bold=True)
    run2 = p.add_run(desc)
    set_run_font(run2, 'Calibri', 11)

doc.add_page_break()

# 2. 产品架构
add_heading(doc, "二、产品架构", 1)

add_heading(doc, "2.1 核心模块", 2)
add_body(doc, "MVP 版本仅包含以下两个核心模块，所有功能均围绕这两个模块展开：")

modules = [
    ("模块一：戒断计时器", "记录用户戒断时长，提供可视化的进度展示和里程碑激励"),
    ("模块二：渴望应对面板", "在渴望高峰时刻，提供 3 秒内可执行的物理打断技术组合"),
]
for title, desc in modules:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    set_run_font(run, 'Calibri', 11, bold=True, color=(30, 58, 95))
    run2 = p.add_run(f" — {desc}")
    set_run_font(run2, 'Calibri', 11)

add_heading(doc, "2.2 数据收集与报告体系", 2)
add_body(doc, "用户数据是产品迭代和形成阶段性报告的基础。MVP 版本的数据收集设计如下：")

add_heading(doc, "（一）数据收集维度", 3)
data_rows = [
    ("戒断记录", "每次戒断开始时间、结束时间（relapse 发生时间）、戒断持续时长", "核心"),
    ("触发事件日志", "触发时间、触发地点（家中/宿舍/其他）、情绪状态（无聊/压力大/孤独/焦虑/其他）、渴望等级（1-10）", "核心"),
    ("渴望应对记录", "每次打开渴望应对面板的时间戳、选择了哪个打断技术、是否完成（做完/放弃）", "核心"),
    ("情绪日志", "每日心情打分（1-5）、简要文字备注（可选）", "基础"),
    ("里程碑成就", "完成的里程碑节点（7天/30天/90天等）、累计戒断天数", "核心"),
    ("周度统计", "每周戒断天数、触发事件次数趋势、渴望应对完成率", "核心"),
]
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t.rows[0].cells
for i, h in enumerate(["数据维度", "具体内容", "优先级"]):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True
for row_data in data_rows:
    row = t.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
doc.add_paragraph()

add_heading(doc, "（二）阶段性报告生成", 3)
add_body(doc, "系统根据本地数据自动生成阶段性报告，供用户查看个人进展：")

report_items = [
    ("每日摘要", "当天戒断状态、情绪打分、是否使用渴望应对、简要评语"),
    ("每周报告", "本周戒断总天数 vs 上周对比、触发事件 TOP3 诱因、渴望应对使用率、本周亮点与建议"),
    ("里程碑庆祝页", "达成 7/14/30/60/90 天时，自动生成可视化庆祝页面（不分享，纯本地）"),
]
for name, desc in report_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"【{name}】")
    set_run_font(run, 'Calibri', 11, bold=True, color=(68, 114, 196))
    run2 = p.add_run(desc)
    set_run_font(run2, 'Calibri', 11)

add_body(doc, "报告以卡片形式呈现于「我的进度」页面，支持滑动浏览。数据不出本地，不上传服务器。", indent=False)

add_heading(doc, "2.3 隐私与伪装设计", 2)
privacy_items = [
    "应用名称建议使用「清流」或其他健康/冥想类词汇，不含任何敏感词",
    "应用图标采用抽象涟漪/山形，不暴露产品意图",
    "应用内可设计「假桌面」功能：打开应用默认显示计算器或天气界面，通过双击顶部特定区域进入真实功能（高级功能）",
    "通知内容完全中性化：使用「健康挑战」、「21天自由计划」等词汇",
    "所有数据存储于本地 SQLite 数据库，可选加密，不强制云端",
    "应用进入需 PIN 码或生物识别（Face ID/指纹）保护",
]
for item in privacy_items:
    add_bullet(doc, item)

doc.add_page_break()

# 3. 功能详细说明
add_heading(doc, "三、核心功能详细说明", 1)

# 3.1 戒断计时器
add_heading(doc, "3.1 戒断计时器（模块一）", 1)

add_heading(doc, "（一）功能概述", 2)
add_body(doc, "戒断计时器是产品的核心数据锚点。它记录用户自设定戒断目标以来的持续时长，以天/小时/分钟为单位的倒计时或正计时形式呈现，配合可视化进度条和里程碑节点，提供持续的成就感。")

add_heading(doc, "（二）用户交互流程", 2)
steps = [
    "首次启动：用户设定个人戒断目标（可选择预设模板如「21天挑战」或自定义天数）",
    "确认后立即开始计时，计时器显示于首页（主屏）最显著位置",
    "用户可随时手动记录一次 relapse（通过首页悬浮按钮，不可错过），系统记录时间戳并重新开始计时",
    "relapse 后不批判，显示「恢复比重新开始更容易」的科学事实卡片，然后平滑重置计时器",
]
for i, s in enumerate(steps, 1):
    add_numbered(doc, s, i, indent=0)

add_heading(doc, "（三）计时器显示规格", 2)

timer_specs = [
    ("主屏大字计时", "以「XX天 XX小时 XX分钟」格式显示，数字醒目（字号 ≥ 32pt）"),
    ("进度条", "围绕计时器的环形进度条，显示距下一个里程碑（如30天）的进度百分比"),
    ("里程碑徽章", "达成里程碑时徽章动画弹出（7天青铜/30天白银/60天黄金/90天钻石等）"),
    ("次级信息", "计时器下方显示：「已避免 XX 次诱惑」「本周第 X 天连续」等补充数据"),
    ("桌面小部件", "Android 支持桌面小部件（Widget），显示天数和简单进度"),
]
for name, spec in timer_specs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"● {name}：")
    set_run_font(run, 'Calibri', 11, bold=True)
    run2 = p.add_run(spec)
    set_run_font(run2, 'Calibri', 11)

add_heading(doc, "（四）relapse 处理逻辑", 2)
add_body(doc, "relapse（戒断中断）是正常过程，不是失败。产品对此的设计原则：")
relapse_items = [
    "relapse 后即时弹出「重置计时器」界面，显示「你的身体平均需要21天建立新习惯，但完全恢复上一次戒断状态只需要更短时间」",
    "询问是否记录 relapse 原因（可选，不强制），用于后续触发分析",
    "relapse 后不显示任何惩罚性文案，不显示连续失败记录",
    "历史 relapse 数据仅用户可见，不做公开排名或对比",
]
for item in relapse_items:
    add_bullet(doc, item)

doc.add_paragraph()

# 3.2 渴望应对面板
add_heading(doc, "3.2 渴望应对面板（模块二）— 产品核心突破口", 1)

add_heading(doc, "（一）设计理念", 2)
add_body(doc, "色情/手淫成瘾最困难的时刻不是「我想戒」，而是「我现在马上就要」。大多数戒断工具把最好的内容藏在菜单深处，让用户在渴望高峰去「浏览功能列表」。", indent=False)
add_body(doc, "渴望应对面板的设计原则：**3秒内让用户做一件真实的物理事件**，强制打断渴求神经回路。", indent=False)

add_heading(doc, "（二）打断技术库（共 11 项，均经神经科学验证）", 2)
add_body(doc, "以下每一项技术均配备：触发说明 + 3 步操作指引 + 预计耗时 + 原理简述（用户点击可见）", indent=False)

techniques = [
    ("T1 冷水激活", "哺乳动物潜水反射，强制迷走神经激活，5秒内接管皮质决策区", "用冷水洗30秒脸，冷水开到最大", "30秒", "迷走神经激活多巴胺清零"),
    ("T2 爆发性运动", "肾上腺素飙升清空多巴胺库存，给爬行脑一个物理出口", "做20个深蹲或原地冲刺15秒，带动作GIF指引", "45秒", "肾上腺素-多巴胺拮抗"),
    ("T3 强薄荷味觉打断", "强烈感官刺激强制重启注意力锚点", "含一口强薄荷糖或用薄荷牙膏漱口", "即时", "三叉神经强烈激活覆盖注意"),
    ("T4 90秒呼吸重启", "延长呼气激活副交感神经，消退渴求高峰（渴求高峰通常持续90秒）", "呼气15秒 × 6轮，自动播放引导节奏（音频/动画）", "90秒", "副交感神经激活降压降心率"),
    ("T5 感官接地54321", "强制重启前额叶：54321感官 grounding", "依次说出：5件看见的东西、4件摸到的东西、3件听见的声音、2件闻到的气味、1件尝到的味道", "60秒", "前额叶重启覆盖自动化渴求"),
    ("T6 冷热交替", "血管收缩扩张节律激活副交感交替", "冷水30秒 → 热水30秒 → 冷水30秒（淋浴情境）", "90秒", "血管舒缩振荡调节自主神经"),
    ("T7 冰敷后颈", "脊髓背根神经冷受体强烈激活，产生镇静效应", "用冰袋或冰水瓶敷在后颈中部，每次30秒，间隔10秒，共3次", "2分钟", "冷受体激活下行抑制通路"),
    ("T8 屏息冲刺", "短暂缺氧触发体内应急保护机制，清空多巴胺池", "深吸气后全力冲刺20米（室内原地冲刺亦可），然后深呼吸3次，重复", "60秒", "缺氧-再氧化-神经递质重置"),
    ("T9 冷水浸泡手腕", "尺侧腕部皮肤冷受体密集，激活迷走神经", "将两侧手腕浸入冷水（盆中加冰块），持续60秒", "60秒", "局部冷刺激迷走神经反射"),
    ("T10 强震感换场", "通过强烈触觉刺激强制转移注意力至体感皮层", "用力捏大腿外侧（疼痛阈值以下）3次，或使用减压玩具（握力球）", "即时", "脊髓上行感觉门控覆盖奖赏回路"),
    ("T11 冷水洗脸+自我对话", "冷激活 + 认知重构同步进行", "冷水洗完脸后，对着镜子说出：「这只是多巴胺在作祟，90秒后我会更好」", "45秒", "冷激活+认知解离协同"),
]

t2 = doc.add_table(rows=1, cols=5)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = t2.rows[0].cells
for i, h in enumerate(["技术名称", "神经原理", "操作步骤", "耗时", "关键机制"]):
    hdr2[i].text = h
    hdr2[i].paragraphs[0].runs[0].bold = True
for row_data in techniques:
    row = t2.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph()

add_heading(doc, "（三）渴望应对面板交互设计", 2)

# 面板布局说明
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run("面板布局（用户点击首页「我现在很想」按钮后展示）：")
set_run_font(run, 'Calibri', 11, bold=True, color=(30, 58, 95))

panel_layout = """
+--------------------------------------------------+
|  [ ⚡ 此刻渴望很正常，你完全可以渡过去 ]           |
|                                                   |
|  [ 🔵 冷水洗30秒脸         ]  <- 最优先推荐位置   |
|  [ ⬜ 20个深蹲打断它       ]                      |
|  [ ⬜ 90秒呼吸重启        ]                      |
|                                                   |
|  [ 📋 写下我现在在想什么  ]  <- 认知重构入口     |
+--------------------------------------------------+
"""
add_body(doc, panel_layout, indent=False)

panel_rules = [
    "每次展示 3 个打断技术按钮，随机抽取，避免形成固定习惯（技术组合每日轮换）",
    "最推荐位置始终是冷水洗脸（T1），物理打断效率最高",
    "「写下我现在在想什么」是认知重构引导入口，点击展开文本输入框，配合引导语：「这个想法是真的吗，还是多巴胺在骗你？」",
    "面板整体设计简洁，禁止任何滚动，必须一眼看完所有选项",
    "技术按钮点击后直接进入对应计时/指引界面，无子菜单",
]
for rule in panel_rules:
    add_bullet(doc, rule)

add_heading(doc, "（四）技术执行引导界面", 2)
add_body(doc, "用户选择某个打断技术后，进入引导界面：")
guide_items = [
    "顶部显示技术名称和原理（1句话，如「冷水激活了你的迷走神经」）",
    "中间显示倒计时器（T1 30秒/T4 90秒等）或步骤指示",
    "底部「完成」按钮，计时结束后点亮",
    "完成后弹出简短庆祝：「你刚刚成功拦截了一次冲动，这让你离目标更近一步」",
    "放弃按钮（不显眼）：用户可随时放弃，但放弃后不批判，只显示「你随时可以再试」",
]
for item in guide_items:
    add_bullet(doc, item)

add_heading(doc, "（五）认知重构日志", 2)
add_body(doc, "「写下我现在在想什么」模块包含以下功能：")
cog_items = [
    "引导语：「现在脑海中出现的画面或想法是什么？」（文本输入，不限字数）",
    "自动生成 AI 辅助分析（本地模型）：「你这个想法中包含了哪些自动化思维？」",
    "提供认知重构选项：「把这个想法换一种说法，你会怎么说？」",
    "日志保存在本地，可回溯查看历史触发思维模式",
]
for item in cog_items:
    add_bullet(doc, item)

add_heading(doc, "（六）渴望应对数据收集", 2)
urge_data = [
    "每次打开面板的时间戳",
    "用户当前渴望等级（1-10，选择面板前须选填）",
    "用户选择的技术编号",
    "是否完成（完成/放弃）",
    "完成耗时",
    "relapse 是否在此次应对后发生（关联数据）",
]
for item in urge_data:
    add_bullet(doc, item)

doc.add_page_break()

# 3.3 触发识别与预测
add_heading(doc, "3.3 触发识别与风险时段感知（本地数据）", 1)

add_heading(doc, "（一）触发日志记录入口", 2)
add_body(doc, "用户在感受到可能引发冲动的情况时，可主动记录触发事件：")
trigger_items = [
    "快速入口：首页悬浮「+」按钮，一键打开触发日志",
    "记录字段：时间（自动）、地点（手动选择）、情绪类型（单选：无聊/压力/孤独/焦虑/看见诱因/睡前/其他）、渴望等级（1-10）",
    "每次记录耗时不超过 15 秒",
]
for item in trigger_items:
    add_bullet(doc, item)

add_heading(doc, "（二）风险时段感知（完全本地算法）", 2)
add_body(doc, "基于本地日志数据，自动识别用户的高风险时空模式：")
risk_items = [
    "连续独处超过 2 小时 → 系统悄悄激活更强内容过滤（可选）",
    "凌晨 1:00-5:00 → 高风险时段，通知变为「健康提醒」（中性）",
    "周中 vs 周末风险差异 → 生成「你的周末风险高于周中」的报告提醒",
    "情绪组合触发：孤独+无聊 组合出现 ≥ 3次/周 → 推荐进入进阶课程",
]
for item in risk_items:
    add_bullet(doc, item)

add_heading(doc, "（三）预测性干预（非骚扰推送）", 2)
add_body(doc, "系统不推送任何暴露意图的消息，采用中性话术：")
predictive = [
    "高风险时段前：推送「今晚感觉怎么样？」（不暴露目的）",
    "3 天未打开 app：推送「你的健康挑战还在进行中」（温和召回）",
    "里程碑临近时：推送「你快达成30天目标了！」（正向激励）",
]
for item in predictive:
    add_bullet(doc, item)

doc.add_page_break()

# 3.4 我的进度与报告
add_heading(doc, "3.4 我的进度与报告页", 1)

add_heading(doc, "（一）页面结构", 2)
progress_items = [
    "顶部：戒断计时器（当前显示）",
    "第一屏卡片：本月数据总览（戒断天数/触发次数/渴望应对完成率）",
    "第二屏卡片：本周 vs 上周对比（柱状图）",
    "第三屏卡片：触发诱因分布（饼图，TOP3 诱因高亮）",
    "第四屏卡片：里程碑进度（进度条 + 徽章墙）",
    "底部：每日摘要入口、每周报告入口",
]
for item in progress_items:
    add_bullet(doc, item)

add_heading(doc, "（二）数据可视化规范", 2)
viz_items = [
    "所有图表基于本地数据生成，不请求网络",
    "图表配色中性柔和（蓝-灰-白），不使用任何敏感色",
    "柱状图/折线图用于趋势，饼图用于诱因分布",
    "里程碑徽章：图标式展示，点击可见达成日期",
]
for item in viz_items:
    add_bullet(doc, item)

doc.add_page_break()

# 4. UI/UX 设计规范
add_heading(doc, "四、UI/UX 设计规范", 1)

add_heading(doc, "4.1 视觉风格", 2)
visual = [
    ("色彩体系", "主色：深蓝 #1F3A5F（信任/专业）；辅色：浅蓝 #4472C4（活力/健康）；强调色：薄荷绿 #4CAF50（清新/正向）；背景：#F5F7FA（柔和）"),
    ("字体", "中文字体：苹方（PingFang）/思源黑体；英文字体：Calibri；无衬线为主"),
    ("圆角", "所有卡片/按钮统一使用 12px 圆角，柔和不刺眼"),
    ("图标", "线性图标为主，风格统一，传达积极正向感"),
    ("间距", "内容区统一使用 16px 安全边距，元素间距 8px/12px/16px 递进"),
]
for name, desc in visual:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"● {name}：")
    set_run_font(run, 'Calibri', 11, bold=True)
    run2 = p.add_run(desc)
    set_run_font(run2, 'Calibri', 11)

add_heading(doc, "4.2 导航结构", 2)
add_body(doc, "底部 Tab 导航，共 3 个入口：")
nav_items = [
    ("首页", "戒断计时器 + 渴望应对快捷入口 + 触发日志入口"),
    ("进度", "我的进度 + 数据报告（每日/每周/里程碑）"),
    ("设置", "PIN码管理/隐私设置/里程碑目标设置/数据导出/关于我们"),
]
for name, desc in nav_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"● {name}：")
    set_run_font(run, 'Calibri', 11, bold=True)
    run2 = p.add_run(desc)
    set_run_font(run2, 'Calibri', 11)

add_heading(doc, "4.3 应用名称与图标", 2)
name_items = [
    "建议应用名：「清流」（简洁，中性，有正态感）",
    "应用图标：抽象涟漪/水滴图形，蓝色系主色，完全去敏感化",
    "不要使用「戒」「自律」「克制」等直接词汇，避免道德审视感",
]
for item in name_items:
    add_bullet(doc, item)

add_heading(doc, "4.4 无障碍与包容性设计", 2)
a11y_items = [
    "支持深色模式（夜间友好，减少屏幕亮度对睡前用户的影响）",
    "字体大小可调节（辅助功能）",
    "所有按钮最小触控区域 44×44pt（iOS HIG / Material Design 标准）",
    "色彩对比度符合 WCAG AA 标准（≥ 4.5:1）",
]
for item in a11y_items:
    add_bullet(doc, item)

doc.add_page_break()

# 5. 技术规格
add_heading(doc, "五、技术规格", 1)

add_heading(doc, "5.1 技术选型", 2)
tech_table = [
    ("跨平台框架", "Flutter（推荐）或 React Native", "一次开发，iOS/Android 双端"),
    ("数据存储", "SQLite（本地）+ 可选 AES-256 加密", "完全离线，数据不出设备"),
    ("图表", "fl_chart（Flutter）或 react-native-chart-kit", "本地渲染，无网络请求"),
    ("通知推送", "Flutter Local Notifications / 本地推送", "预测性干预通知，不依赖服务端"),
    ("内容过滤（可选）", "Android：本地 VPN DNS 过滤；iOS：Screen Time API", "可选项，非强制"),
    ("应用壳/伪装", "flutter_secure_storage + 条件启动", "支持双重应用形态"),
    ("隐私保护", "应用 PIN 码 / Face ID / 指纹", "每次打开应用需验证"),
    ("AI 认知辅助（可选）", "ONNX 本地小模型（端侧推理）", "离线可用，保护隐私"),
]
t3 = doc.add_table(rows=1, cols=3)
t3.style = 'Table Grid'
hdr3 = t3.rows[0].cells
for i, h in enumerate(["技术项", "选型方案", "备注"]):
    hdr3[i].text = h
    hdr3[i].paragraphs[0].runs[0].bold = True
for row_data in tech_table:
    row = t3.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
doc.add_paragraph()

add_heading(doc, "5.2 数据安全", 2)
security_items = [
    "所有用户数据存储于设备本地，不上传任何服务器",
    "relapse 日志、触发日志、渴望应对记录均为私密数据",
    "可设置应用 PIN 码或生物识别锁",
    "如提供数据导出功能（JSON格式），仅导出至用户本地，不支持云端",
    "应用内无任何广告，无第三方追踪 SDK",
]
for item in security_items:
    add_bullet(doc, item)

add_heading(doc, "5.3 性能要求", 2)
perf_items = [
    "应用冷启动时间 < 2 秒",
    "渴望应对面板打开时间 < 1 秒（3秒内完成目标的一部分）",
    "数据库查询时间 < 100ms",
    "应用安装包大小 < 50MB（Flutter Release Build）",
    "后台运行内存占用 < 100MB",
]
for item in perf_items:
    add_bullet(doc, item)

doc.add_page_break()

# 6. 里程碑与迭代计划
add_heading(doc, "六、MVP 里程碑与产品迭代计划", 1)

add_heading(doc, "6.1 MVP 交付范围（v1.0）", 2)
mvp_items = [
    "戒断计时器（核心计时 + 里程碑 + relapse 处理）",
    "渴望应对面板（11项打断技术 + 执行引导 + 认知日志）",
    "触发日志（快速记录入口 + 本地分析）",
    "我的进度（数据总览 + 每周报告 + 可视化图表）",
    "隐私保护（PIN码/生物识别 + 应用中性化）",
    "预测性干预（非骚扰推送）",
]
for item in mvp_items:
    add_bullet(doc, item)

add_heading(doc, "6.2 第二版迭代方向（v1.1-1.2）", 2)
iter2 = [
    "匿名社区/小组打卡功能（需严格设计正向互助机制）",
    "AI 教练（基于本地日志的个性化建议推送）",
    "双重应用形态（计算器伪装入口）",
    "内容过滤增强（DNS 层 adult 域名拦截）",
    "数据云端加密同步（可选，会员功能）",
]
for item in iter2:
    add_bullet(doc, item)

add_heading(doc, "6.3 第三版迭代方向（v2.0+）", 2)
iter3 = [
    "家长知情模式（针对未成年用户）",
    "双人共同模式（伴侣/朋友相互支持）",
    "专业版：对接心理咨询资源（危机热线 + 线上咨询入口）",
    "Web 端数据看板（用户自行查看导出数据）",
]
for item in iter3:
    add_bullet(doc, item)

doc.add_page_break()

# 7. 风险与合规
add_heading(doc, "七、风险与合规", 1)

add_heading(doc, "7.1 合规要求", 2)
compliance_items = [
    "中国内地：《互联网信息服务管理办法》合规，不存储/传播敏感内容",
    "美国：COPPA 合规（如涉及13岁以下用户，禁止数据收集）",
    "欧盟：GDPR 合规（如在欧盟地区提供服务），提供数据删除权利",
    "应用商店审核：确保应用名称/图标/描述完全去敏感化，避免被拒",
    "心理健康内容：明确免责声明（非专业心理咨询工具），接入正规心理危机热线",
]
for item in compliance_items:
    add_bullet(doc, item)

add_heading(doc, "7.2 已知风险", 2)
risks = [
    ("应用商店审查风险", "高", "色情成瘾类应用可能被误判为成人内容，需准备完整的应用说明文档"),
    ("青少年用户引导风险", "中", "需明确产品不针对未成年人进行医学建议，可提供家长知情模式"),
    ("relapse 后心理危机", "低", "需内置心理危机热线（非紧急按钮），relapse 后不批判"),
    ("技术依赖风险", "中", "产品需逐步引导用户建立内在自律能力，而非永远依赖外部工具"),
]
t4 = doc.add_table(rows=1, cols=3)
t4.style = 'Table Grid'
hdr4 = t4.rows[0].cells
for i, h in enumerate(["风险项", "等级", "应对措施"]):
    hdr4[i].text = h
    hdr4[i].paragraphs[0].runs[0].bold = True
for row_data in risks:
    row = t4.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
doc.add_paragraph()

doc.add_page_break()

# 8. 附录
add_heading(doc, "附录 A：渴望应对技术神经科学原理速查", 1)

add_body(doc, "本文档中的 11 项打断技术均基于以下神经科学机制：")
appendix_table = [
    ("多巴胺拮抗", "T2爆发性运动", "肾上腺素与多巴胺共享神经通路，肾上腺素激增可短暂清空多巴胺储备"),
    ("迷走神经激活", "T1冷水/T6冷热交替/T9冰敷手腕", "冷水激活哺乳动物潜水反射，副交感神经激活，降低心率和血压，压制渴求"),
    ("前额叶重启", "T5感官接地54321", "强制将注意力转移至体感/视觉皮层，覆盖自动化奖赏回路的神经活动"),
    ("三叉神经覆盖", "T3强薄荷味觉", "强感官刺激激活三叉神经，产生强烈存在感信号，覆盖奖赏渴望信号"),
    ("呼吸-副交感耦合", "T4 90秒呼吸", "延长呼气（15秒）激活迷走神经，迷走神经激活后抑制多巴胺释放"),
    ("体感门控", "T10强震感换场", "强烈触觉激活脊髓上行门控，产生体感覆盖效应，阻断奖赏回路信号"),
    ("缺氧-再氧化神经递质重置", "T8屏息冲刺", "短暂缺氧触发应急机制，再氧化过程重置神经递质平衡"),
]
t5 = doc.add_table(rows=1, cols=3)
t5.style = 'Table Grid'
hdr5 = t5.rows[0].cells
for i, h in enumerate(["神经机制", "对应技术", "科学原理说明"]):
    hdr5[i].text = h
    hdr5[i].paragraphs[0].runs[0].bold = True
for row_data in appendix_table:
    row = t5.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
doc.add_paragraph()

add_heading(doc, "附录 B：推荐产品名称备选", 1)
names = [
    "清流 — 简洁，中性，有正态感，最推荐",
    "FreeM / 自由模式 — 强调自由感，不暴露意图",
    "涟漪 — 抽象自然，有健康感，不敏感",
    "清朗 — 清新明朗，暗示清晰的生活状态",
]
for name in names:
    add_bullet(doc, name)

# ===== 保存 =====
output_path = "/Users/mac/.mavis/sessions/mvs_eff1face857e48f1b1ed0b7db5986352/workspace/清流_产品需求文档 PRD_v1.0.docx"
doc.save(output_path)
print(f"文档已保存至：{output_path}")